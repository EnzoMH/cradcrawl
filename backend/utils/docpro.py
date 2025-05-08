# docpro.py - 간소화된 버전
import io
import os
import tempfile
import logging
from typing import Optional, Dict, Any
from fastapi import UploadFile, HTTPException

# 문서 처리 라이브러리
from PyPDF2 import PdfReader
from langchain_teddynote.document_loaders import HWPLoader
import pandas as pd

# 선택적 라이브러리 (설치된 경우에만 사용)
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    
# Excel 처리 라이브러리 추가
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

# 로깅 설정
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)
if not logger.handlers:
    handler = logging.StreamHandler()
    handler.setFormatter(logging.Formatter('[%(asctime)s] %(levelname)s: %(message)s'))
    logger.addHandler(handler)

async def process_file(file: UploadFile) -> str:
    """
    파일을 처리하여 텍스트 추출
    
    Args:
        file: 업로드된 파일
        
    Returns:
        str: 추출된 텍스트
    """
    try:
        # 파일 확장자 확인
        ext = os.path.splitext(file.filename.lower())[1][1:]
        
        # 파일 내용 읽기
        content = await file.read()
        if not content:
            raise HTTPException(status_code=400, detail="빈 파일입니다")
        
        # 파일 형식에 따라 처리
        if ext in ['pdf']:
            return process_pdf(content)
        elif ext in ['hwp', 'hwpx']:
            return process_hwp(content)
        elif ext == 'docx':
            return process_docx(content)
        elif ext == 'doc':
            return process_doc(content)
        elif ext in ['xlsx', 'xls']:
            return process_excel(content)
        else:
            raise HTTPException(status_code=400, detail=f"지원하지 않는 파일 형식: {ext}")
            
    except Exception as e:
        logger.error(f"파일 처리 오류: {str(e)}")
        raise HTTPException(status_code=500, detail=f"파일 처리 중 오류 발생: {str(e)}")

def process_pdf(content: bytes) -> str:
    """PDF 파일 처리"""
    logger.info("PDF 처리 시작")
    text = ""
    
    # 1. PyMuPDF 시도 (더 좋은 결과를 제공)
    if PYMUPDF_AVAILABLE:
        try:
            pdf_stream = io.BytesIO(content)
            doc = fitz.open(stream=pdf_stream, filetype="pdf")
            
            pages_text = []
            for page in doc:
                pages_text.append(page.get_text())
                
            text = "\n".join(pages_text)
            doc.close()
            
            if text.strip():
                return text
        except Exception as e:
            logger.warning(f"PyMuPDF 처리 실패: {str(e)}")
    
    # 2. PyPDF2 사용
    try:
        pdf_stream = io.BytesIO(content)
        reader = PdfReader(pdf_stream)
        
        text_parts = []
        for page in reader.pages:
            text_parts.append(page.extract_text())
            
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF 처리 실패: {str(e)}")
        raise

def process_hwp(content: bytes) -> str:
    """HWP/HWPX 파일 처리"""
    logger.info("HWP 파일 처리 시작")
    
    # 임시 파일 생성
    with tempfile.NamedTemporaryFile(suffix='.hwp', delete=False) as temp_file:
        temp_path = temp_file.name
        temp_file.write(content)
    
    try:
        # HWPLoader를 사용하여 텍스트 추출
        hwp_loader = HWPLoader(file_path=temp_path)
        docs = hwp_loader.load()
        
        if not docs:
            raise ValueError("HWP 파일에서 텍스트를 추출할 수 없습니다")
        
        return docs[0].page_content
    except Exception as e:
        logger.error(f"HWP 처리 실패: {str(e)}")
        raise
    finally:
        # 임시 파일 삭제
        try:
            os.remove(temp_path)
        except Exception as e:
            logger.warning(f"임시 파일 삭제 실패: {str(e)}")

def process_docx(content: bytes) -> str:
    """DOCX 파일 처리"""
    logger.info("DOCX 처리 시작")
    
    if not DOCX_AVAILABLE:
        raise ImportError("DOCX 처리를 위한 python-docx 라이브러리가 필요합니다")
    
    # 임시 파일 생성
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        temp_path = temp_file.name
        temp_file.write(content)
    
    try:
        # python-docx를 사용하여 텍스트 추출
        doc = docx.Document(temp_path)
        
        # 텍스트 추출 (단락 + 표)
        text_parts = []
        
        # 단락 추출
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # 표 추출
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return '\n'.join(text_parts)
    except Exception as e:
        logger.error(f"DOCX 처리 실패: {str(e)}")
        raise
    finally:
        # 임시 파일 삭제
        try:
            os.remove(temp_path)
        except Exception as e:
            logger.warning(f"임시 파일 삭제 실패: {str(e)}")

def process_doc(content: bytes) -> str:
    """DOC 파일 처리 (제한적 지원)"""
    logger.warning("DOC 파일 형식은 제한적으로 지원됩니다")
    
    # 임시 파일 생성
    with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
        temp_path = temp_file.name
        temp_file.write(content)
    
    try:
        # 외부 명령어 실행 (antiword가 설치된 경우)
        import subprocess
        try:
            result = subprocess.run(
                ['antiword', temp_path], 
                capture_output=True, 
                text=True, 
                check=True
            )
            return result.stdout
        except (subprocess.SubprocessError, FileNotFoundError):
            logger.warning("antiword를 사용할 수 없습니다. 텍스트 추출이 제한적일 수 있습니다.")
            
            # 대체 방법: 바이너리 데이터에서 텍스트 추출 시도
            text_bytes = b''
            for i in range(0, len(content) - 1):
                if content[i] >= 32 and content[i] < 127 and content[i+1] == 0:
                    text_bytes += bytes([content[i]])
            
            # 다양한 인코딩 시도
            for encoding in ['utf-8', 'cp949', 'euc-kr', 'latin1']:
                try:
                    return text_bytes.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("DOC 파일에서 텍스트를 추출할 수 없습니다")
    except Exception as e:
        logger.error(f"DOC 처리 실패: {str(e)}")
        raise
    finally:
        # 임시 파일 삭제
        try:
            os.remove(temp_path)
        except Exception as e:
            logger.warning(f"임시 파일 삭제 실패: {str(e)}")

def process_excel(content: bytes) -> str:
    """Excel 파일(XLSX, XLS) 처리"""
    logger.info("Excel 파일 처리 시작")
    
    if not PANDAS_AVAILABLE:
        raise ImportError("Excel 처리를 위한 pandas 라이브러리가 필요합니다")
    
    try:
        # BytesIO를 사용하여 메모리에서 읽기
        excel_stream = io.BytesIO(content)
        
        # 모든 시트 데이터 추출
        dfs = pd.read_excel(excel_stream, sheet_name=None, engine='openpyxl')
        
        # 모든 시트의 데이터를 텍스트로 변환
        text_parts = []
        
        for sheet_name, df in dfs.items():
            # 시트 이름 추가
            text_parts.append(f"[시트: {sheet_name}]")
            
            # 행 추가 (컬럼 이름 포함)
            headers = " | ".join(str(col) for col in df.columns)
            text_parts.append(headers)
            
            # 구분선 추가
            text_parts.append("-" * len(headers))
            
            # 데이터 행 추가
            for _, row in df.iterrows():
                row_text = " | ".join(str(val) if not pd.isna(val) else "" for val in row)
                text_parts.append(row_text)
            
            # 시트 구분을 위한 빈 줄
            text_parts.append("\n")
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Excel 처리 실패: {str(e)}")
        raise

def clean_text(text: str) -> str:
    """추출된 텍스트 정리"""
    if not text:
        return ""
    
    # 연속된 공백 제거
    import re
    text = re.sub(r'\s+', ' ', text)
    
    # 빈 줄 정리
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text.strip()

